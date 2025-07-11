import os
from docx import Document

def replace_placeholder_in_para(paragraph, context):
    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, val in context.items():
        new_text = new_text.replace(f"[{key}]", str(val))
    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

def generate_word(template_path, context, filename):
    doc = Document(template_path)

    # Replace placeholders in paragraphs
    for p in doc.paragraphs:
        replace_placeholder_in_para(p, context)

    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_in_para(p, context)

    # Special Case: Insert table for Exam NOC under [PFNumber]
    if context.get("LetterType") == "Exam NOC":
        for i, paragraph in enumerate(doc.paragraphs):
            if "[PFNumber]" in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "PF Number"
                hdr[1].text = "Employee Name"
                hdr[2].text = "Designation"
                hdr[3].text = "NOC Year"
                hdr[4].text = "Application No."
                hdr[5].text = "Exam Name"

                row = table.add_row().cells
                row[0].text = str(context["PFNumberVal"])
                row[1].text = context["EmployeeName"]
                row[2].text = context["Designation"]
                row[3].text = str(context["NOCYear"])
                row[4].text = str(context["AppNo"])
                row[5].text = context["ExamName"]
                break

    # Save generated file
    output_path = os.path.join("generated_letters", filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path