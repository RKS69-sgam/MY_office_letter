import streamlit as st
import pandas as pd
import datetime
import os
from docx import Document

# === Utility Function ===
def generate_word(template_path, context, filename):
    doc = Document(template_path)

    def replace_placeholder_in_para(paragraph, context):
        full_text = ''.join(run.text for run in paragraph.runs)
        new_text = full_text
        for key, val in context.items():
            new_text = new_text.replace(f"[{key}]", str(val))
        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ''
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    for p in doc.paragraphs:
        replace_placeholder_in_para(p, context)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_in_para(p, context)

    output_path = os.path.join("generated_letters", filename)
    os.makedirs("generated_letters", exist_ok=True)
    doc.save(output_path)
    return output_path

def download_word(path):
    with open(path, "rb") as f:
        st.download_button(
            label="Download Letter",
            data=f,
            file_name=os.path.basename(path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# === Load Data ===
class_file = "assets/CLASS-III EMPLOYEES.xlsx"
class_df = pd.read_excel(class_file, sheet_name="Sheet1")
class_df["Display"] = class_df.apply(
    lambda r: f"{r['PF No.']} - {r['HRMS ID']} - {r.name+1} - {r['Employee Name']}", axis=1
)

# === Streamlit UI ===
st.title("OFFICE OF THE SSE/PW/SGAM")
letter_type = st.selectbox("Select Letter Type", ["Engine Pass Letter", "Card Pass Letter"])

if letter_type in ["Engine Pass Letter", "Card Pass Letter"]:
    st.subheader(letter_type)
    selected_emp = st.selectbox("Select Employee", class_df["Display"])
    letter_date = st.date_input("Letter Date", value=datetime.date.today())

    selected_row = class_df[class_df["Display"] == selected_emp].iloc[0]

    context = {
        "EmployeeName": selected_row["Employee Name"],
        "Designation": selected_row["Designation"],
        "PFNumber": selected_row["PF No."],
        "LetterDate": letter_date.strftime("%d-%m-%Y"),
    }

    dor_val = selected_row.get("DOR", None)
    if pd.notnull(dor_val):
        context["DOR"] = pd.to_datetime(dor_val).strftime("%d-%m-%Y")
    else:
        context["DOR"] = ""

    if st.button("Generate Letter"):
        if letter_type == "Engine Pass Letter":
            template_path = "assets/Engine Pass letter temp.docx"
            save_name = f"EnginePass-{context['EmployeeName'].strip()}.docx"
            col_to_update = "Engine Pass Renewal Application Date"
        else:
            template_path = "assets/Card Pass letter temp.docx"
            save_name = f"CardPass-{context['EmployeeName'].strip()}.docx"
            col_to_update = "Card Pass Renewal Application Date"

        word_path = generate_word(template_path, context, save_name)
        st.success("Letter generated successfully.")
        download_word(word_path)

        row_index = class_df[class_df["Display"] == selected_emp].index[0]
        class_df.at[row_index, col_to_update] = letter_date.strftime("%d-%m-%Y")
        class_df.drop(columns=["Display"], inplace=True, errors="ignore")
        class_df.to_excel(class_file, sheet_name="Sheet1", index=False)
        st.success("Register updated.")