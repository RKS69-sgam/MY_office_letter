import streamlit as st
import pandas as pd
import os
import base64
from docx import Document
from datetime import date, timedelta
import datetime
from docx.shared import Inches
from dateutil.relativedelta import relativedelta
from pathlib import Path

# --- Configuration and Data Loading ---
BASE_DIR = Path(__file__).parent
OUTPUT_FOLDER = BASE_DIR / "generated_letters"
OUTPUT_FOLDER.mkdir(exist_ok=True)

# File paths
template_files = {
    "Duty Letter (For Absent)": "assets/Absent Duty letter temp.docx",
    "SF-11 For Other Reason": "assets/SF-11 temp.docx",
    "Sick Memo": "assets/SICK MEMO temp..docx",
    "General Letter": "assets/General Letter temp.docx",
    "Exam NOC": "assets/Exam NOC Letter temp.docx",
    "SF-11 Punishment Order": "assets/SF-11 Punishment order temp.docx",
    "Quarter Allotment Letter": "assets/Quarter Allotment temp.docx",
    "Update Employee Database": None,
    "Engine Pass Letter": "assets/Engine Pass letter temp.docx",
    "Card Pass Letter": "assets/Card Pass letter temp.docx",
    "PME Memo": "assets/pme_memo_temp.docx"
}

# --- Robust Global DataFrames Loading ---
def safe_load_df(excel_path, csv_path=None, sheet_name=None, is_master=False):
    """Attempts to load DataFrame from Excel, then CSV if Excel fails. Returns empty DataFrame on full failure."""
    if csv_path is None:
        csv_path = excel_path.replace(".xlsx", ".xlsx - Sheet1.csv")

    if is_master:
        # Special handling for employee_master, which is a dictionary of DataFrames
        try:
            return pd.read_excel(excel_path, sheet_name=None)
        except Exception:
            # st.warning(f"Failed to load Excel master data ({excel_path}). Trying CSV fallback...")
            try:
                # Assuming one main sheet for the master CSV
                return {"Apr.25": pd.read_csv(csv_path)}
            except Exception:
                # st.error(f"Failed to load Employee Master from CSV: {csv_e}")
                return {"Apr.25": pd.DataFrame()}
    else:
        # Standard DataFrame loading
        try:
            return pd.read_excel(excel_path, sheet_name=sheet_name)
        except Exception:
            # st.warning(f"Failed to load Excel data ({excel_path}). Trying CSV fallback...")
            try:
                if Path(csv_path).exists():
                     return pd.read_csv(csv_path)
                else:
                    # If CSV file path is complicated and doesn't exist, return empty DF
                    return pd.DataFrame()
            except Exception:
                # st.error(f"Failed to load data from CSV: {csv_e}")
                return pd.DataFrame() # Always return a DataFrame on failure


quarter_df = safe_load_df(
    excel_path="assets/QUARTER REGISTER.xlsx",
    csv_path="assets/QUARTER REGISTER.xlsx - Sheet1.csv",
    sheet_name="Sheet1"
)

employee_master = safe_load_df(
    excel_path="assets/EMPLOYEE MASTER DATA.xlsx",
    csv_path="assets/EMPLOYEE MASTER DATA.xlsx - Apr.25.csv",
    is_master=True
)

sf11_register_path = "assets/SF-11 Register.xlsx"
sf11_register = safe_load_df(
    excel_path=sf11_register_path,
    csv_path="assets/SF-11 Register.xlsx - SSE-SGAM.csv",
    sheet_name="SSE-SGAM"
)

noc_register_path = "assets/Exam NOC_Report.xlsx"
df_noc = safe_load_df(
    excel_path=noc_register_path,
    csv_path="assets/Exam NOC_Report.xlsx - Sheet1.csv",
    sheet_name=None 
)

# Robust initialization check for df_noc (Fix for recurring error)
# Ensures df_noc is a DataFrame before checking its attributes
if not isinstance(df_noc, pd.DataFrame) or df_noc.empty or "PF Number" not in df_noc.columns:
    df_noc = pd.DataFrame(columns=["PF Number", "Employee Name", "Designation", "NOC Year", "Application No.", "Exam Name"])
    
# --- End of Robust Global DataFrames Loading ---


# --- Helper Functions ---
def replace_placeholder_in_para(paragraph, context):
    full_text = ''.join(run.text for run in paragraph.runs)
    new_text = full_text
    for key, val in context.items():
        val_str = str(val) if val is not None else ""
        new_text = new_text.replace(f"[{key}]", val_str)
        new_text = new_text.replace(f"{{{{ {key} }}}}", val_str)
        new_text = new_text.replace(f"{{{{ {key}}}}}", val_str)
        new_text = new_text.replace(f"{{{{{key}}}}}", val_str)

    if new_text != full_text:
        for run in paragraph.runs:
            run.text = ''
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

# === Generate Word Function ===
def generate_word(template_path, context, filename):
    try:
        doc = Document(template_path)
    except Exception as e:
        st.error(f"Error loading template {template_path}: {e}")
        return None
        
    # Replace in paragraphs
    for p in doc.paragraphs:
        replace_placeholder_in_para(p, context)
    # Replace in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholder_in_para(p, context)

    # ✅ Exam NOC Table Insertion (Updated for individual inputs and accurate columns)
    if context.get("LetterType") == "Exam NOC" and context.get("EmployeeData"):
        for i, paragraph in enumerate(doc.paragraphs):
            if "[PFNumber]" in paragraph.text:
                # Remove placeholder paragraph
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
                
                # Insert table
                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                table.autofit = True
                
                # Table Headers
                hdr = table.rows[0].cells
                hdr[0].text = "Sr. No."
                hdr[1].text = "PF Number"
                hdr[2].text = "Employee Name"
                hdr[3].text = "Designation"
                hdr[4].text = "Exam's Name" 
                hdr[5].text = "Term of NOC" 
                
                # Add rows for each employee
                for idx, emp_data in enumerate(context["EmployeeData"]):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(idx + 1)
                    row_cells[1].text = str(emp_data["PF Number"])
                    row_cells[2].text = emp_data["Employee Name"]
                    row_cells[3].text = emp_data["Designation"]
                    row_cells[4].text = emp_data["Exam Name"] 
                    row_cells[5].text = emp_data["Term of NOC"] 
                    
                break
    
    output_path = os.path.join("generated_letters", filename)
    doc.save(output_path)
    return output_path

# === Download Function ===
def download_word(path):
    with open(path, "rb") as f:
        b64 = base64.b64encode(f.read()).decode()
        name = os.path.basename(path)
        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{name}">Download Word File</a>'
        st.markdown(href, unsafe_allow_html=True)

# Placeholder for handle_engine_card_pass function if it's in another file
def handle_engine_card_pass(letter_type):
    st.warning(f"The logic for '{letter_type}' is a placeholder. It needs to be implemented or imported from `engine_card_pass.py`.")
    if letter_type in ["Engine Pass Letter", "Card Pass Letter"]:
        return "N/A"
    return ""

def parse_date_safe(date_str):
    """Attempts to parse date string into a datetime.date object. Returns None on failure."""
    if pd.isna(date_str):
        return None
    try:
        if isinstance(date_str, datetime.date):
            return date_str
        if isinstance(date_str, pd.Timestamp):
            return date_str.date()
        
        formats = ["%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%Y-%m-%d %H:%M:%S"]
        
        for fmt in formats:
            try:
                return datetime.datetime.strptime(str(date_str).split()[0], fmt).date()
            except ValueError:
                continue
    except Exception:
        pass
    return None

def format_date_safe(date_val):
    """Safely converts date objects or strings to DD-MM-YYYY format."""
    parsed_date = parse_date_safe(date_val)
    if parsed_date:
        return parsed_date.strftime("%d-%m-%Y")
    return "N/A" if pd.isna(date_val) else str(date_val)

def get_age_service_length(dob_date, doa_date):
    """Calculates age and service length from DOB and DOA date objects."""
    today = date.today()
    age, service_year, service_month = 0, 0, 0

    if dob_date:
        age_delta = relativedelta(today, dob_date)
        age = age_delta.years

    if doa_date:
        service_delta = relativedelta(today, doa_date)
        service_year = service_delta.years
        service_month = service_delta.months

    return age, service_year, service_month


# === PME MEMO UI FUNCTION ===
def render_pme_memo_ui(row):
    st.markdown("#### आवधिक चिकित्सा परीक्षा (PME) मेमो विवरण")
    
    pf_number = str(row.get("PF No.", "NO_PF"))
    dob_date = parse_date_safe(row.get("DOB"))
    doa_date = parse_date_safe(row.get("DOA"))
    
    initial_last_pme = parse_date_safe(row.get("LAST PME")) or date.today()
    initial_pme_due = parse_date_safe(row.get("PME DUE")) or date.today()
    initial_med_cat = str(row.get("Medical category", "A3"))
    
    father_name = str(row.get("FATHER'S NAME", "N/A"))
    designation_hindi = row.get("Designation in Hindi") if pd.notna(row.get("Designation in Hindi")) else row.get("DESIGNATION", "N/A")

    age, service_year, service_month = get_age_service_length(dob_date, doa_date)
    
    # --- UI Inputs (Editable) ---
    st.subheader("मेडिकल/ड्यूटी डिटेल्स (आवश्यक)")
    
    last_pme_date = st.date_input("पिछली PME दिनांक (Last PME Date)", value=initial_last_pme, key=f"pme_last_date_{pf_number}")
    pme_due_date = st.date_input("PME ड्यू दिनांक (PME Due Date)", value=initial_pme_due, key=f"pme_due_date_{pf_number}")
    med_cat = st.text_input("मेडिकल कैटेगरी (Medical Category)", value=initial_med_cat, key=f"pme_med_cat_{pf_number}")

    st.subheader("अन्य मेमो विवरण")
    last_place = st.text_input("पिछली परीक्षा का स्थान (Last Exam Place)", value="ACMS/NKJ", key=f"pme_last_place_{pf_number}")
    examiner = st.text_input("डॉक्टर का पदनाम (Examiner Designation)", value="ACMS", key=f"pme_examiner_{pf_number}")
    first_mark = st.text_input("शारीरिक पहचान चिन्ह 1 (Physical Mark 1)", value="A mole on the left hand.", key=f"pme_mark1_{pf_number}")
    second_mark = st.text_input("शारीरिक पहचान चिन्ह 2 (Physical Mark 2)", value="A scar on the right elbow.", key=f"pme_mark2_{pf_number}")
    
    # --- Context Formatting ---
    
    return {
        "dob": format_date_safe(dob_date),
        "doa": format_date_safe(doa_date),
        "name": row.get("Employee Name", ""),
        "age": age,
        "father_name": father_name,
        "designation": designation_hindi, 
        "medical_category": med_cat,
        "last_examined_date": format_date_safe(last_pme_date),
        "last_place": last_place,
        "examiner": examiner,
        "service_year": service_year,
        "service_month": service_month,
        "first_physical_mark": first_mark,
        "second_physical_mark": second_mark,
        "current_date": date.today().strftime("%d-%m-%Y"),
        "LetterType": "PME Memo"
    }

# === Register Update Function ===
def update_registers(letter_type, context, letter_date, pf, hname, desg, patra_kr=None, noc_employees=None, q_selected=None):
    """Handles all logic for updating the global DataFrames and saving them to file."""
    global sf11_register, df_noc, quarter_df 

    # --- SF-11 Register Entry (For Other Reason or Duty Letter) ---
    if letter_type in ["SF-11 For Other Reason", "Duty Letter (For Absent)"]:
        new_entry = pd.DataFrame([{
            "पी.एफ. क्रमांक": pf, "कर्मचारी का नाम": hname, "पदनाम": desg, "पत्र क्र.": context["LetterNo"],
            "दिनांक": letter_date.strftime("%d-%m-%Y"), "दण्ड का विवरण": context["Memo"]
        }])
        sf11_register = pd.concat([sf11_register, new_entry], ignore_index=True)
        sf11_register.to_excel(sf11_register_path, sheet_name="SSE-SGAM", index=False)
        st.success("SF-11 register updated.")

    # --- SF-11 Register Update (For Punishment) ---
    if letter_type == "SF-11 Punishment Order" and patra_kr is not None:
        mask = (sf11_register["पी.एफ. क्रमांक"] == pf) & (sf11_register["पत्र क्र."] == patra_kr)
        if mask.any():
            i = sf11_register[mask].index[0]
            sf11_register.loc[i, "दण्डादेश क्रमांक"] = context["LetterNo"]
            sf11_register.loc[i, "दण्ड का विवरण"] = context["Memo"]
            sf11_register.loc[i, "पावती का दिनांक"] = context["pawati_date"].strftime("%d-%m-%Y")
            sf11_register.loc[i, "यदि प्रत्‍युत्तर प्राप्‍त हुआ हो तो दिनांक"] = context["pratyuttar_date"].strftime("%d-%m-%Y")
            sf11_register.to_excel(sf11_register_path, sheet_name="SSE-SGAM", index=False)
            st.success("SF-11 register updated.")
        else:
            st.warning("चयनित कर्मचारी के लिए पत्र क्रमांक के आधार पर प्रविष्टि नहीं मिली।")
            
    # --- Quarter Allotment Register Update ---
    if letter_type == "Quarter Allotment Letter" and q_selected is not None:
        # Check if 'Display' column exists before dropping it (used for selection UI)
        if 'Display' in quarter_df.columns:
            temp_quarter_df = quarter_df.copy()
            temp_quarter_df["Display"] = temp_quarter_df.apply(lambda r: f"{r['STATION']} - {r['QUARTER NO.']}", axis=1)
            i = temp_quarter_df[temp_quarter_df["Display"] == q_selected].index[0]
        else:
            # Fallback (shouldn't happen if loading is correct)
            i = quarter_df[quarter_df['QUARTER NO.'] == q_selected.split(" - ")[1]].index[0]


        quarter_df.loc[i, "PF No."] = pf
        quarter_df.loc[i, "EMPLOYEE NAME"] = hname
        quarter_df.loc[i, "OCCUPIED DATE"] = letter_date.strftime("%d-%m-%Y")
        quarter_df.loc[i, "STATUS"] = "OCCUPIED"
        
        quarter_df.to_excel(quarter_file, sheet_name="Sheet1", index=False)
        st.success("Quarter Register updated.")


    # === Exam NOC Register Entry (Updated for individual inputs) ===
    if letter_type == "Exam NOC" and noc_employees is not None and noc_employees:
        new_noc_entries = []
        year = date.today().year
        
        for emp in noc_employees:
            pf_num = emp["PF Number"]
            exam_name = emp["Exam Name"]
            
            # Check existing count for the current employee/year
            df_match = df_noc[(df_noc["PF Number"] == pf_num) & (df_noc["NOC Year"] == year)]
            current_count = df_match.shape[0]

            if current_count < 4:
                new_noc_entries.append({
                    "PF Number": pf_num,
                    "Employee Name": emp["Employee Name"],
                    "Designation": emp["Designation"],
                    "NOC Year": year,
                    "Application No.": current_count + 1,
                    "Exam Name": exam_name
                })
            else:
                pass 

        if new_noc_entries:
            df_noc = pd.concat([df_noc, pd.DataFrame(new_noc_entries)], ignore_index=True)
            df_noc.to_excel(noc_register_path, index=False)
            st.success(f"{len(new_noc_entries)} Exam NOC entries added to register.")


# === UI ===
st.title("OFFICE OF THE SSE/PW/SGAM")

password = st.text_input("Enter Password", type="password")
if password == "sgam@4321":
    st.success("Access Granted!")

    letter_type = st.selectbox("Select Letter Type", list(template_files.keys()))

    master_df = employee_master["Apr.25"]
    master_df["Display"] = master_df.apply(lambda r: f"{r['PF No.']} - {r['Employee Name']} - {r['UNIT / MUSTER NUMBER']} - {r['DESIGNATION']}", axis=1)

    dor_str = ""
    pf = hname = desg = unit_full = unit = short = letter_no = ""
    row = None
    letter_date = date.today()
    patra_kr = None
    q_selected = None
    noc_employees = [] # List to hold data for selected employees for NOC

    # --- Conditional Employee Selection & Data Fetching ---

    if letter_type == "Exam NOC":
        selected_display_names = st.multiselect("Select Employees for NOC", master_df["Display"].dropna())
        
        if selected_display_names:
            selected_rows = master_df[master_df["Display"].isin(selected_display_names)]
            
            st.markdown("### परीक्षा विवरण (Exam Details - Individual)")
            
            for idx, r in selected_rows.iterrows():
                pf_num = str(r["PF No."])
                employee_name = r["Employee Name in Hindi"] if pd.notna(r["Employee Name in Hindi"]) else r["Employee Name"]
                desg_val = r["Designation in Hindi"] if pd.notna(r["Designation in Hindi"]) else r["DESIGNATION"]
                
                # Check for existing NOC limit (optional check, main check is in update_registers)
                year = date.today().year
                df_match = df_noc[(df_noc["PF Number"] == pf_num) & (df_noc["NOC Year"] == year)]
                current_count = df_match.shape[0]

                if current_count >= 4:
                    st.error(f"**{employee_name} ({pf_num})**: इस वर्ष पहले ही 4 NOC ले चुका है। इसे छोड़ दिया जाएगा।")
                    continue

                st.markdown(f"**{employee_name} ({pf_num})**")
                
                # Individual Inputs for Exam Name and Term
                exam_name = st.text_input(f"Exam Name (Current Count: {current_count})", key=f"exam_name_{pf_num}", 
                                          placeholder="Enter Exam Name")
                term = st.text_input(f"Term of NOC", key=f"noc_term_{pf_num}", 
                                     placeholder="e.g., 2024-25")
                
                if exam_name and term:
                    noc_employees.append({
                        "PF Number": pf_num,
                        "Employee Name": employee_name,
                        "Designation": desg_val,
                        "Exam Name": exam_name,
                        "Term of NOC": term
                    })
            
            # Use the data of the first selected employee for the general letter context
            if not selected_rows.empty:
                r = selected_rows.iloc[0]
                pf = r["PF No."]
                hname = r["Employee Name in Hindi"] if pd.notna(r["Employee Name in Hindi"]) else r["Employee Name"]
                desg = r["Designation in Hindi"] if pd.notna(r["Designation in Hindi"]) else r["DESIGNATION"]
                unit_full = str(r["UNIT / MUSTER NUMBER"])
                unit = unit_full[:2]
                short = r["SF-11 short name"] if pd.notna(r["SF-11 short name"]) else "STF"
                letter_no = f"{short}/{unit}/NOC"
        
        letter_date = st.date_input("Letter Date", value=date.today())
        
    elif letter_type == "SF-11 Punishment Order":
        # Single select for Punishment Order
        df = sf11_register
        df["Display"] = df.apply(lambda r: f"{r['पी.एफ. क्रमांक']} - {r['कर्मचारी का नाम']} - {r['पत्र क्र.']} - {r['दिनांक']}", axis=1)
        selected = st.selectbox("Select Employee", df["Display"].dropna())
        row = df[df["Display"] == selected].iloc[0]
        patra_kr = row["पत्र क्र."]
        dandadesh_krmank = f"{patra_kr}/D-1"
        pf = row["पी.एफ. क्रमांक"]
        hname = row["कर्मचारी का नाम"]
        desg = row.get("पदनाम", "")
        unit_full = patra_kr.split("/", 1)[1]
        unit = unit_full[-7:]
        short = patra_kr.split("/")[0]
        letter_no = dandadesh_krmank
        sf11date = row["दिनांक"]
        letter_date = st.date_input("Letter Date", value=date.today())
        
    elif letter_type == "General Letter" or letter_type == "Update Employee Database":
        letter_date = st.date_input("Letter Date", value=date.today())
        
    elif letter_type in ["Engine Pass Letter", "Card Pass Letter"]:
        dor_str = handle_engine_card_pass(letter_type)
        selected = st.selectbox("Select Employee", master_df["Display"].dropna())
        row = master_df[master_df["Display"] == selected].iloc[0]
        pf = row["PF No."]
        hname = row["Employee Name in Hindi"] if pd.notna(row["Employee Name in Hindi"]) else row["Employee Name"]
        desg = row["Designation in Hindi"] if pd.notna(row["Designation in Hindi"]) else row["DESIGNATION"]
        unit_full = str(row["UNIT / MUSTER NUMBER"])
        unit = unit_full[:2]
        short = row["SF-11 short name"] if pd.notna(row["SF-11 short name"]) else "STF"
        letter_no = f"{short}/{unit}/{unit}"
        letter_date = st.date_input("Letter Date", value=date.today())
        
    else: # Default single select for Duty Letter, Sick Memo, PME Memo, Quarter Allotment
        selected = st.selectbox("Select Employee", master_df["Display"].dropna())
        row = master_df[master_df["Display"] == selected].iloc[0]
        pf = row["PF No."]
        hname = row["Employee Name in Hindi"] if pd.notna(row["Employee Name in Hindi"]) else row["Employee Name"]
        desg = row["Designation in Hindi"] if pd.notna(row["Designation in Hindi"]) else row["DESIGNATION"]
        unit_full = str(row["UNIT / MUSTER NUMBER"])
        unit = unit_full[:2]
        short = row["SF-11 short name"] if pd.notna(row["SF-11 short name"]) else "STF"
        letter_no = f"{short}/{unit}/{unit}"
        letter_date = st.date_input("Letter Date", value=date.today())


    # === Common context for template replacement ===
    context = {
        "LetterDate": letter_date.strftime("%d-%m-%Y"),
        "EmployeeName": hname,
        "Designation": desg,
        "PFNumber": pf,
        "ShortName": short,
        "Unit": unit,
        "UnitNumber": unit,
        "LetterNo": letter_no,
        "DutyDate": "", "FromDate": "", "ToDate": "", "JoinDate": "", "Memo": "",
        "OfficerUnit": "", "Subject": "", "Reference": "", "CopyTo": "", "DOR": dor_str
    }
    
    # --- Letter Specific UI ---
    if letter_type == "Duty Letter (For Absent)":
        mode = st.selectbox("Mode", ["SF-11 & Duty Letter Only", "Duty Letter Only"])
        fd = st.date_input("From Date")
        td = st.date_input("To Date", value=date.today())
        jd = st.date_input("Join Date", value=td + timedelta(days=1))
        context.update({
            "FromDate": fd.strftime("%d-%m-%Y"),
            "ToDate": td.strftime("%d-%m-%Y"),
            "JoinDate": jd.strftime("%d-%m-%Y"),
            "DutyDate": jd.strftime("%d-%m-%Y"),
            "Memo": f"आप बिना किसी पूर्व सूचना के दिनांक {fd.strftime('%d-%m-%Y')} से {td.strftime('%d-%m-%Y')} तक कुल {(td-fd).days+1} दिवस कार्य से अनुपस्थित थे, जो कि रेल सेवक होने के नाते आपकी रेल सेवा निष्ठा के प्रति घोर लापरवाही को प्रदर्शित करता है। अतः आप कामों व भूलो के फेहरिस्त धारा 1, 2 एवं 3 के उल्लंघन के दोषी पाए जाते है।"
        })
    elif letter_type == "SF-11 For Other Reason":
        memo_input = st.text_area("Memo")
        context["Memo"] = memo_input + " जो कि रेल सेवक होने नाते आपकी रेल सेवा निष्ठा के प्रति घोर लापरवाही को प्रदर्शित करता है। अतः आप कामों व भूलो के फेहरिस्त धारा 1, 2 एवं 3 के उल्लंघन के दोषी पाए जाते है।"
    elif letter_type == "General Letter":
        context["FileName"] = st.selectbox("File Name", ["", "STAFF-IV", "OFFICE ORDER", "STAFF-III", "QAURTER-1", "ARREAR", "CEA/STAFF-IV", "CEA/STAFF-III", "PW-SGAM", "MISC."])
        officer_option = st.selectbox("अधिकारी/कर्मचारी", ["", "सहायक मण्‍डल अभियंता", "मण्‍डल अभिंयता (पूर्व)","मुख्‍य चिकित्‍सा अधीक्षक", "मण्‍डल अभिंयता (पश्चिम)", "मण्‍डल रेल प्रबंधक (कार्मिक)", "मण्‍डल रेल प्रबंधक (कार्य)", "वरिष्‍ठ खण्‍ड अभियंता (रेल पथ)", "वरिष्‍ठ खण्‍ड अभियंता (कार्य)", "वरिष्‍ठ खण्‍ड अभियंता (विद्युत)", "वरिष्‍ठ खण्‍ड अभियंता (T&D)", "वरिष्‍ठ खण्‍ड अभियंता (S&T)", "वरिष्‍ठ खण्‍ड अभियंता (USFD)", "वरिष्‍ठ खण्‍ड अभियंता (PW/STORE)", "कनिष्‍ठ अभियंता (रेल पथ)", "कनिष्‍ठ अभियंता (कार्य)", "कनिष्‍ठ अभियंता (विद्युत)", "कनिष्‍ठ अभियंता (T&D)", "कनिष्‍ठ अभियंता (S&T)", "शाखा सचिव (WCRMS)", "मण्‍डल अध्‍यक्ष (WCRMS)", "मण्‍डल सचिव (WCRMS)", "महामंत्री (WCRMS)", "अन्‍य"])
        if officer_option == "अन्‍य": officer_option = st.text_input("अन्‍य का नाम/पदनाम/एजेंसी का नाम लिखें")
        context["OfficerName"] = officer_option
        
        beyohari_officers = ["सहायक मण्‍डल अभियंता", "वरिष्‍ठ खण्‍ड अभियंता (कार्य)", "वरिष्‍ठ खण्‍ड अभियंता (विद्युत)", "वरिष्‍ठ खण्‍ड अभियंता (T&D)", "वरिष्‍ठ खण्‍ड अभियंता (S&T)", "शाखा सचिव (WCRMS)"]
        jbp_officers = ["मण्‍डल अभिंयता (पूर्व)", "मुख्‍य चिकित्‍सा अधीक्षक", "मण्‍डल अभिंयता (पश्चिम)", "मण्‍डल रेल प्रबंधक (कार्मिक)", "मण्‍डल रेल प्रबंधक (कार्य)", "वरिष्‍ठ खण्‍ड अभियंता (S&T)", "वरिष्‍ठ खण्‍ड अभियंता (USFD)", "वरिष्‍ठ खण्‍ड अभियंता (PW/STORE)", "मण्‍डल अध्‍यक्ष (WCRMS)", "मण्‍डल सचिव (WCRMS)", "महामंत्री (WCRMS)"]
        
        if officer_option == "कनिष्‍ठ अभियंता (रेल पथ)": address_choices = ["निवासरोड", "भरसेड़ी", "गजराबहरा", "गोंदवाली", "अन्‍य"]
        elif officer_option in beyohari_officers: address_choices = ["प.म.रे. ब्‍योहारी", "अन्‍य"]
        elif officer_option in jbp_officers: address_choices = ["प.म.रे. जबलपुर", "अन्‍य"]
        else: address_choices = ["", "प.म.रे. ब्‍योहारी", "प.म.रे. जबलपुर", "सरईग्राम", "देवराग्राम", "बरगवॉं", " निवासरोड", "भरसेड़ी", "गजराबहरा", "गोंदवाली", "अन्‍य"]
        
        address_option = st.selectbox("पता", address_choices)
        if address_option == "अन्‍य": address_option = st.text_input("अन्‍य का पता लिखें")
        context["OfficeAddress"] = address_option
        
        context["Subject"] = f"विषय:- {st.text_input('विषय')}" if st.text_input('विषय').strip() else ""
        context["Reference"] = f"संदर्भ:- {st.text_input('संदर्भ')}" if st.text_input('संदर्भ').strip() else ""
        context["Memo"] = st.text_area("मुख्‍य विवरण")
        copy_input = st.text_input("प्रतिलिपि")
        context["CopyTo"] = f"प्रतिलिपि:- " + "\n".join(
            [c.strip() for c in copy_input.split(",") if c.strip()]
        ) if copy_input.strip() else ""

    elif letter_type == "Exam NOC":
        if noc_employees:
            # Additional context specific to NOC processing is added here.
            context.update({
                "ExamName": noc_employees[0]["Exam Name"], # Using first employee data for top-level context (for filename/general placeholders)
                "Term": noc_employees[0]["Term of NOC"],
                "NOCYear": date.today().year,
                "LetterType": "Exam NOC",
                "EmployeeData": noc_employees
            })

    elif letter_type == "SF-11 Punishment Order" and row is not None:
        st.markdown("#### SF-11 Register से विवरण")
        st.markdown(f"**आरोप का विवरण:** {row.get('आरोप का विवरण', '—')}")
        pawati_date = st.date_input("पावती का दिनांक", value=date.today())
        pratyuttar_date = st.date_input("यदि प्रत्‍युत्तर प्राप्‍त हुआ हो तो दिनांक", value=date.today())
        
        context["Memo"] = st.selectbox("Punishment Type", [
            "आगामी देय एक वर्ष की वेतन वृद्धि असंचयी प्रभाव से रोके जाने के अर्थदंड से दंडित किया जाता है।",
            "आगामी देय एक वर्ष की वेतन वृद्धि संचयी प्रभाव से रोके जाने के अर्थदंड से दंडित किया जाता है।",
            "आगामी देय एक सेट सुविधा पास तत्काल प्रभाव से रोके जाने के दंड से दंडित किया जाता है।",
            "आगामी देय एक सेट PTO तत्काल प्रभाव से रोके जाने के दंड से दंडित किया जाता है।",
            "आगामी देय दो सेट सुविधा पास तत्काल प्रभाव से रोके जाने के दंड से दंडित किया जाता है।",
            "आगामी देय दो सेट PTO तत्काल प्रभाव से रोके जाने के दंड से दंडित किया जाता."
        ])
        context["Dandadesh"] = letter_no
        context["LetterNo."] = patra_kr
        context["Unit"] = unit
        context["SF-11Date"] = format_date_safe(sf11date)
        context["pawati_date"] = pawati_date
        context["pratyuttar_date"] = pratyuttar_date
        
    elif letter_type == "Quarter Allotment Letter" and row is not None:
        # NOTE: Using try-except block here because quarter_df may not have all columns initially
        try:
            quarter_df["Display"] = quarter_df.apply(lambda r: f"{r['STATION']} - {r['QUARTER NO.']}", axis=1)
            q_selected = st.selectbox("Select Quarter", quarter_df["Display"].dropna())
            qrow = quarter_df[quarter_df["Display"] == q_selected].iloc[0]
            station = qrow["STATION"]
            qno = qrow["QUARTER NO."]
            context.update({
                "QuarterNo.": qno,
                "Station": station,
                "q_selected": q_selected
            })
        except Exception as e:
            st.warning(f"Quarter selection failed: {e}. Ensure QUARTER REGISTER data integrity.")
            q_selected = None


    elif letter_type == "PME Memo" and row is not None:
        pme_context = render_pme_memo_ui(row)
        context.update(pme_context)

    elif letter_type == "Update Employee Database":
        st.subheader("Update Employee Database")
        st.info("The actual Update Employee Database logic is complex and needs to be fully implemented here, followed by saving the master Excel file.")
        
    
    # Generate letter command
    if st.button("Generate Letter"):

        if letter_type == "Update Employee Database":
            st.info("Employee Database update is handled by the dedicated UI section above. No letter generated.")
        elif row is None and letter_type not in ["General Letter", "Exam NOC"]:
            st.error("Please select an employee before generating the letter.")
        elif letter_type == "Exam NOC" and not noc_employees:
            st.error("Please select at least one employee for the Exam NOC and fill in the details.")
        else:
            # --- Generate the Document ---
            word_path = None
            if letter_type == "Duty Letter (For Absent)" and 'mode' in locals() and mode == "SF-11 & Duty Letter Only":
                duty_path = generate_word(template_files["Duty Letter (For Absent)"], context, f"DutyLetter-{hname}.docx")
                sf11_path = generate_word(template_files["SF-11 For Other Reason"], context, f"SF-11-{hname}.docx")
                if duty_path: download_word(duty_path)
                if sf11_path: download_word(sf11_path)
                word_path = duty_path 
            elif letter_type == "General Letter":
                today_str = datetime.datetime.now().strftime("%d-%m-%Y")
                filename_parts = [context.get("FileName", "").replace("/", "-").strip(), context.get("OfficerName", "").strip(), today_str, context.get("Subject", "").replace("विषय:-", "").strip()[:10]]
                final_name = " - ".join(part for part in filename_parts if part).replace(" - -", "").strip()
                word_path = generate_word(template_files["General Letter"], context, f"{final_name}.docx")
                if word_path: download_word(word_path)
            elif letter_type == "PME Memo":
                filename = f"PME_Memo-{context['EmployeeName'].strip()}-{context['dob']}.docx"
                word_path = generate_word(template_files["PME Memo"], context, filename)
                if word_path: download_word(word_path)
                st.success("PME Memo generated successfully.")
            elif letter_type == "Quarter Allotment Letter":
                filename = f"QuarterAllotmentLetter-{hname}.docx"
                word_path = generate_word(template_files["Quarter Allotment Letter"], context, filename)
                if word_path: download_word(word_path)
            elif letter_type == "Exam NOC":
                filename = f"ExamNOC_Multi-{context['EmployeeData'][0]['PF Number']}-{len(context['EmployeeData'])}.docx"
                word_path = generate_word(template_files["Exam NOC"], context, filename)
                if word_path: download_word(word_path)
                st.success(f"Multi-Employee Exam NOC generated successfully for {len(context['EmployeeData'])} employees.")
            else:
                filename = f"{letter_type.replace('/', '-')}-{hname}.docx"
                word_path = generate_word(template_files[letter_type], context, filename)
                if word_path: download_word(word_path)

            
            # --- Update Registers only if document generation was successful ---
            if word_path:
                update_registers(
                    letter_type=letter_type, 
                    context=context, 
                    letter_date=letter_date, 
                    pf=pf, 
                    hname=hname, 
                    desg=desg, 
                    patra_kr=patra_kr, 
                    noc_employees=noc_employees if letter_type == "Exam NOC" else None, 
                    q_selected=context.get("q_selected")
                )


elif password != "":
    st.error("Incorrect Password")
